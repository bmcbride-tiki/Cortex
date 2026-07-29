# =============================================================================
# test_import_from_word.py
# -----------------------------------------------------------------------------
# Copyright 2025 Brian McBride at Tiki-1 Studio
# WHAT THIS FILE DOES
#   Checks that import_from_word.py's `run()` extracts real paragraph text
#   from a real generated `.docx`, and fails cleanly on a missing file.
#
# WHAT IT INTERACTS WITH
#   - `import_from_word.py`, the file under test.
#   - `python-docx`, used here only to generate a throwaway source `.docx`.
# =============================================================================

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from import_from_word import ImportFromWord


def test_reads_paragraphs():
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document
        src = Path(tmp) / "in.docx"
        doc = Document()
        doc.add_paragraph("Hello")
        doc.add_paragraph("World")
        doc.save(str(src))

        result = ImportFromWord().run(str(src))
        assert result["success"] is True
        assert result["text"] == "Hello\nWorld"


def test_missing_file_fails_cleanly():
    result = ImportFromWord().run("/no/such/file.docx")
    assert result["success"] is False


if __name__ == "__main__":
    test_reads_paragraphs()
    test_missing_file_fails_cleanly()
    print("All import_from_word self-checks passed.")
