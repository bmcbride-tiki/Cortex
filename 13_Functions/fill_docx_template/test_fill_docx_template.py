# =============================================================================
# test_fill_docx_template.py
# -----------------------------------------------------------------------------
# WHAT THIS FILE DOES
#   Checks that fill_docx_template.py's `run()` replaces a `{{ content }}`
#   token in a real generated .docx template, and fails cleanly on a
#   missing template file.
#
# WHAT IT INTERACTS WITH
#   - `fill_docx_template.py`, the file under test.
#   - `python-docx`, used here to generate a throwaway template and read
#     back the filled result.
# =============================================================================

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from fill_docx_template import FillDocxTemplate


def test_fills_content_token():
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document
        template_path = Path(tmp) / "template.docx"
        doc = Document()
        doc.add_paragraph("Report body: {{ content }}")
        doc.save(str(template_path))

        result = FillDocxTemplate().run(
            content_text="the filled-in text",
            template_path=str(template_path),
            output_dir=tmp,
        )
        assert result["success"] is True
        out_doc = Document(result["file_path"])
        assert out_doc.paragraphs[0].text == "Report body: the filled-in text"


def test_missing_template_fails_cleanly():
    result = FillDocxTemplate().run("text", "/no/such/template.docx", "")
    assert result["success"] is False


if __name__ == "__main__":
    test_fills_content_token()
    test_missing_template_fails_cleanly()
    print("All fill_docx_template self-checks passed.")
