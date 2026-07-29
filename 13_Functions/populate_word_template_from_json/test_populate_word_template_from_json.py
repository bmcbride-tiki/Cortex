# =============================================================================
# test_populate_word_template_from_json.py
# -----------------------------------------------------------------------------
# Copyright 2025 Brian McBride at Tiki-1 Studio
# WHAT THIS FILE DOES
#   Checks that populate_word_template_from_json.py's `run()` fills
#   multiple distinct named placeholders in a real generated `.docx`
#   template, and fails cleanly on a missing template or empty data.
#
# WHAT IT INTERACTS WITH
#   - `populate_word_template_from_json.py`, the file under test.
#   - `python-docx`, used here to generate a throwaway template and read
#     back the filled result.
# =============================================================================

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from populate_word_template_from_json import PopulateWordTemplateFromJson


def test_fills_multiple_named_placeholders():
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document
        template_path = Path(tmp) / "template.docx"
        doc = Document()
        doc.add_paragraph("Name: {{ name }}")
        doc.add_paragraph("Trade: {{trade}}")
        doc.save(str(template_path))

        result = PopulateWordTemplateFromJson().run(
            data={"name": "Alex", "trade": "Electrician"},
            template_path=str(template_path),
            output_dir=tmp,
        )
        assert result["success"] is True
        out_doc = Document(result["file_path"])
        assert out_doc.paragraphs[0].text == "Name: Alex"
        assert out_doc.paragraphs[1].text == "Trade: Electrician"


def test_missing_template_fails_cleanly():
    result = PopulateWordTemplateFromJson().run({"a": "b"}, "/no/such/template.docx", "")
    assert result["success"] is False


def test_empty_data_fails_cleanly():
    result = PopulateWordTemplateFromJson().run({}, "/some/template.docx", "")
    assert result["success"] is False


if __name__ == "__main__":
    test_fills_multiple_named_placeholders()
    test_missing_template_fails_cleanly()
    test_empty_data_fails_cleanly()
    print("All populate_word_template_from_json self-checks passed.")
