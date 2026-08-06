# =============================================================================
# import_documents.py
# -----------------------------------------------------------------------------
# Copyright 2025 Brian McBride at Tiki-1 Studio
# WHAT THIS FILE DOES
#   A zero-input Task: picks up every pending .docx/.pdf file sitting in
#   `01_inbox/documents/`, pulls out simple "Label: Value" metadata (title,
#   date, trade) from the document's own text (falling back to parsing the
#   filename if a field isn't found in the body), archives the original
#   file into `02_vault/documents/`, and indexes its full text for search.
#   Runs in two phases: extract everything to a JSONL scratch file first,
#   then load that JSONL into the database -- so a crash partway through
#   loading doesn't lose the already-extracted text.
#
# WHAT IT INTERACTS WITH
#   - `01_inbox/documents/`, read then cleared (originals deleted once
#     safely copied into the vault).
#   - `02_vault/documents/`, where archived originals end up.
#   - `00_System/database.py`'s `get_db_connection()`, writing to the
#     `documents_metadata` table and the `global_search_index` FTS5 table.
#   - `python-docx` (.docx) / `pypdf` (.pdf), for extracting body text.
#   - `test_import_documents.py`, this file's paired test (covers the
#     label-extraction and filename-fallback logic only, not the full
#     file-system/DB run).
# =============================================================================

import json
import re
import sys
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

CURRENT_DIR = Path(__file__).resolve().parent
PARENT_DIR = CURRENT_DIR.parents[1]
CORE_DIR = PARENT_DIR / "00_System"

if str(CORE_DIR) not in sys.path:
    sys.path.append(str(CORE_DIR))

from database import get_db_connection

# Matches lines like "Trade: Electrical" or "Report Date: 2026-01-15" anywhere
# in the document body -- no fixed row/column position assumed, so slightly
# different report layouts all feed the same extractor.
LABEL_LINE_PATTERN = re.compile(r"^([A-Za-z][A-Za-z0-9 /_-]{1,40}):\s+(.+)$")

TITLE_KEYS = ["title", "report title", "document title"]
DATE_KEYS = ["date", "report date", "doc date", "document date"]
TRADE_KEYS = ["trade", "trades"]

SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


class DocumentsImporter:
    def __init__(self):
        self.root_dir = PARENT_DIR
        self.inbox_dir = self._resolve_dir_case_insensitive(self.root_dir / "01_inbox" / "documents")
        self.vault_dir = self.root_dir / "02_vault" / "documents"

    def _resolve_dir_case_insensitive(self, target_path: Path) -> Path:
        """Helper to find the physical path even if it's named 01_Inbox instead of 01_inbox."""
        if target_path.exists():
            return target_path

        parent = target_path.parent
        if parent.exists():
            for child in parent.iterdir():
                if child.name.lower() == target_path.name.lower():
                    return child
        return target_path

    def _read_docx(self, file_path: Path) -> str:
        """Extracts text from both paragraphs and tables inside a DOCX file."""
        try:
            import docx
        except ImportError:
            raise ImportError("The 'python-docx' library is missing. Please run 'pip install python-docx' in your terminal.")

        doc = docx.Document(str(file_path))
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                seen_cells = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and (not seen_cells or seen_cells[-1] != text):
                        seen_cells.append(text)

                if seen_cells:
                    full_text.append(" | ".join(seen_cells))

        return "\n".join(full_text)

    def _read_pdf(self, file_path: Path) -> str:
        """Extracts raw page text from a PDF file (no table-cell structure)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("The 'pypdf' library is missing. Please run 'pip install pypdf' in your terminal.")

        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(p.strip() for p in pages if p.strip())

    def _extract_labeled_fields(self, body_text: str) -> Dict[str, str]:
        """Scans every line for 'Label: Value' pairs and returns them as a flat dict."""
        fields: Dict[str, str] = {}
        for line in body_text.splitlines():
            match = LABEL_LINE_PATTERN.match(line.strip())
            if match:
                label = match.group(1).strip()
                value = match.group(2).strip()
                if label and value:
                    fields[label] = value
        return fields

    def _lookup_field(self, fields: Dict[str, str], keys: List[str]) -> str:
        """Case-insensitive lookup of the first matching key from a candidate list."""
        lowered = {k.lower(): v for k, v in fields.items()}
        for key in keys:
            if key in lowered:
                return lowered[key]
        return ""

    def _parse_filename_fallback(self, filename: str) -> Dict[str, str]:
        """
        Fallback metadata when a field wasn't found in the document body.
        Uses the same YYYY-MM-DD_Title.[ext] convention as import_transcripts.
        """
        stem = Path(filename).stem
        parts = stem.split("_", 1)

        fallback = {
            "date": datetime.now().strftime("%Y-%m-%d"),
            "title": stem.replace("-", " ").strip(),
        }

        if len(parts) >= 1:
            try:
                datetime.strptime(parts[0], "%Y-%m-%d")
                fallback["date"] = parts[0]
                if len(parts) == 2:
                    fallback["title"] = parts[1].replace("-", " ").strip()
            except ValueError:
                pass

        return fallback

    def _split_trades(self, raw_value: str) -> List[str]:
        """Splits a 'Trade' field value on commas/semicolons into a clean list."""
        if not raw_value:
            return []
        parts = re.split(r"[,;]", raw_value)
        return [p.strip() for p in parts if p.strip()]

    def _extract_record(self, file_path: Path) -> Dict[str, Any]:
        """Reads one file and returns its extracted record as a plain dict (JSONL-ready)."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            body_text = self._read_pdf(file_path)
        else:
            body_text = self._read_docx(file_path)

        fields = self._extract_labeled_fields(body_text)
        fallback = self._parse_filename_fallback(file_path.name)

        title = self._lookup_field(fields, TITLE_KEYS) or fallback["title"]
        doc_date = self._lookup_field(fields, DATE_KEYS) or fallback["date"]
        trades = self._split_trades(self._lookup_field(fields, TRADE_KEYS))

        return {
            "title": title,
            "doc_date": doc_date,
            "trades": trades,
            "source_filename": file_path.name,
            "extracted_fields": fields,
            "body_text": body_text,
        }

    def extract_to_jsonl(self, files: List[Path]) -> Path:
        """Phase 1: pull structured info out of every pending file and write it as JSONL."""
        self.vault_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        jsonl_path = self.vault_dir / f"_extracted_{timestamp}.jsonl"

        with open(jsonl_path, "w", encoding="utf-8") as jsonl_file:
            for file_path in files:
                record = self._extract_record(file_path)
                record["_source_path"] = str(file_path)
                jsonl_file.write(json.dumps(record, ensure_ascii=False) + "\n")

        return jsonl_path

    def load_jsonl_to_db(self, jsonl_path: Path) -> Dict[str, Any]:
        """Phase 2: read the JSONL produced above and load each record into cortex.db."""
        report = {"success": True, "files_processed": 0, "records_loaded": 0, "errors": []}

        conn = get_db_connection()
        cursor = conn.cursor()

        with open(jsonl_path, "r", encoding="utf-8") as jsonl_file:
            lines = [line for line in jsonl_file if line.strip()]

        for line in lines:
            record = json.loads(line)
            source_path = Path(record["_source_path"])

            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                vault_filename = f"{source_path.stem}_{timestamp}{source_path.suffix}"
                vault_path = self.vault_dir / vault_filename
                shutil.copy2(str(source_path), str(vault_path))

                cursor.execute("""
                    INSERT INTO documents_metadata (
                        title, doc_date, trades, source_filename, file_path, extracted_fields
                    ) VALUES (?, ?, ?, ?, ?, ?);
                """, (
                    record["title"],
                    record["doc_date"],
                    json.dumps(record["trades"], ensure_ascii=False),
                    record["source_filename"],
                    str(vault_path),
                    json.dumps(record["extracted_fields"], ensure_ascii=False),
                ))
                origin_id = cursor.lastrowid

                cursor.execute("""
                    INSERT INTO global_search_index (origin_id, content_type, body_text)
                    VALUES (?, 'document', ?);
                """, (origin_id, record["body_text"]))

                conn.commit()
                source_path.unlink()

                report["files_processed"] += 1
                report["records_loaded"] += 1

            except Exception as e:
                conn.rollback()
                report["errors"].append(f"Error loading {record.get('source_filename', '?')}: {str(e)}")
                report["success"] = False

        conn.close()
        return report

    def run(self) -> Dict[str, Any]:
        self.inbox_dir.mkdir(parents=True, exist_ok=True)

        files = [
            f for f in self.inbox_dir.iterdir()
            if f.suffix.lower() in SUPPORTED_EXTENSIONS and not f.name.startswith("~$")
        ]

        if not files:
            return {
                "success": False,
                "files_processed": 0,
                "records_loaded": 0,
                "errors": [f"No pending document files (.pdf or .docx) found inside: {self.inbox_dir}"],
            }

        jsonl_path = self.extract_to_jsonl(files)
        report = self.load_jsonl_to_db(jsonl_path)
        report["jsonl_path"] = str(jsonl_path)
        return report


if __name__ == "__main__":
    importer = DocumentsImporter()
    print(importer.run())
