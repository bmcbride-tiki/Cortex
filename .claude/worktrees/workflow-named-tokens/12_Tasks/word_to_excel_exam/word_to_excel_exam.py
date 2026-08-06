"""
---
module_id: 'TRANSFORM-EXAM'
title: 'Word-to-Excel Exam Data Transformer'
tags: ['automation', 'parser', 'docx', 'xlsx']
schema_version: '1.0.0'
---
Self-contained, UI-free engine for converting Word exam banks into Excel ingestion layouts.

Copyright 2025 Brian McBride at Tiki-1 Studio

WHAT THIS FILE DOES
    Parses a Word exam question bank into structured question records
    (stem, four options, correct answer, difficulty, rationale) and writes
    them into a copy of the item-import Excel template. Supports three
    distinct source document formats via separate parsers, selected by the
    caller (`format_type`): "ai" (AI Generated & Validated, heading-based),
    "ai2" (bold-run-based "Rob" format), and "ce" (nested-table CE format).
    Entirely self-contained -- no adapter/API dependency, just python-docx
    + openpyxl.

WHAT IT INTERACTS WITH
    - The input Word question bank .docx, supplied as a CLI argument.
    - `item_import_template.xlsx`, this folder's locked template asset
      (copied, never modified in place).
    - `output/`, the default output folder for the generated workbook (a
      `<user>-Question Importer - <date>.xlsx` file, collision-suffixed).
    - `core_router.py`/`server.py`'s upload endpoint
      (`/api/uploads/word-to-excel-exam`), which saves an uploaded file
      into this folder's `incoming/` and dispatches this script with the
      chosen `format_type`.
"""

import re
import getpass
import shutil
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, List, Dict
from docx import Document
from docx.table import Table
import openpyxl

ANSWER_LETTER_TO_NUMBER = {
    "A": "1",
    "B": "2",
    "C": "3",
    "D": "4",
}

@dataclass
class DocumentLine:
    text: str
    is_heading: bool = False

@dataclass
class BoldDocumentLine:
    text: str
    has_bold_text: bool = False

@dataclass
class QuestionRecord:
    source_type: str
    source_document: str
    item_number: str
    question_text: str
    question_title: str = ""
    option_text_1: str = ""
    option_text_2: str = ""
    option_text_3: str = ""
    option_text_4: str = ""
    option_correct: str = ""
    explanation: str = ""
    question_type: str = "Multiple Choice"
    difficulty: str = ""

class QuestionImporterEngine:
    def __init__(self, template_path: Optional[Path] = None, output_dir: Optional[Path] = None):
        # Establish self-contained localized paths inside the tool's specific subfolder
        self.current_tool_dir = Path(__file__).resolve().parent
        
        self.template_path = template_path or (
            self.current_tool_dir / "item_import_template.xlsx"
        )
        self.output_dir = output_dir or (self.current_tool_dir / "output")
        self.parsed_questions: List[QuestionRecord] = []

    def log(self, message: str):
        """Standardized console logger bypassing missing UI logging text boxes."""
        print(f"[EXAM-TRANSFORMER] {message}", flush=True)

    def execute_pipeline(self, source_file_path: Path, format_type: str) -> Optional[Path]:
        """Orchestrates the complete validation, parsing, and excel compilation pipeline."""
        self.parsed_questions = []
        source_path = Path(source_file_path)

        if not source_path.exists():
            raise FileNotFoundError(f"Source file could not be located at: {source_path}")
        if not self.template_path.exists():
            raise FileNotFoundError(f"Locked template asset missing at: {self.template_path}\n"
                                    f"Please ensure 'item_import_template.xlsx' is inside the tool folder.")

        self.log(f"Beginning analysis of: {source_path.name} using format rule: {format_type}")

        # Route to explicit functional sub-parser types
        if format_type.lower() == "ce":
            self.parsed_questions = self.parse_ce_formatted_file(source_path)
        elif format_type.lower() == "ai2":
            self.parsed_questions = self.parse_ai2_rob_file(source_path)
        else:
            self.parsed_questions = self.parse_ai_validated_file(source_path)

        if not self.parsed_questions:
            self.log("Parsing terminated: Zero complete question matrices were discovered.")
            return None

        self.log(f"Extraction successful: Identified {len(self.parsed_questions)} structural question nodes.")
        
        # Compile records directly to localized spreadsheet outputs
        output_file = self.write_records_to_excel(self.parsed_questions)
        return output_file

    # ==================================================================
    # Format Option 3: AI-2 (Rob) Functional Parser
    # ==================================================================
    def parse_ai2_rob_file(self, file_path: Path) -> List[QuestionRecord]:
        lines = self._extract_bold_document_lines(file_path)
        records: List[QuestionRecord] = []
        current_record: Optional[QuestionRecord] = None
        next_item_number = 1

        for line in lines:
            text = line.text.strip()
            if not text or self._should_skip_ai2_line(text):
                continue

            option_match = self._match_option_line(text)
            if option_match and current_record is not None:
                self._apply_option_to_record(current_record, option_match)
                if line.has_bold_text:
                    answer_letter = option_match.group("letter").strip().upper()
                    current_record.option_correct = ANSWER_LETTER_TO_NUMBER.get(answer_letter, "")

                if self._record_has_core_content(current_record):
                    records.append(current_record)
                    current_record = None
                    next_item_number += 1
                continue

            if option_match and current_record is None:
                continue

            if current_record is not None:
                if self._record_has_any_option(current_record):
                    records.append(current_record)
                    next_item_number += 1
                else:
                    self.log(f"Skipping broken entry at item {current_record.item_number}: Missing options A-D.")

            current_record = QuestionRecord(
                source_type="AI-2 (Rob)",
                source_document=file_path.name,
                item_number=str(next_item_number),
                question_text=text
            )

        if current_record is not None:
            records.append(current_record)

        return self._filter_complete_records(records)

    def _extract_bold_document_lines(self, file_path: Path) -> List[BoldDocumentLine]:
        document = Document(str(file_path))
        lines: List[BoldDocumentLine] = []

        for paragraph in document.paragraphs:
            current_text_parts: List[str] = []
            current_has_bold_text = False

            for run in paragraph.runs:
                run_text = run.text
                if not run_text:
                    continue

                split_parts = run_text.split("\n")
                for index, part in enumerate(split_parts):
                    if part:
                        current_text_parts.append(part)
                        if run.bold is True and part.strip():
                            current_has_bold_text = True

                    if index < len(split_parts) - 1:
                        self._append_bold_line_if_useful(lines, current_text_parts, current_has_bold_text)
                        current_text_parts = []
                        current_has_bold_text = False

            self._append_bold_line_if_useful(lines, current_text_parts, current_has_bold_text)
        return lines

    def _append_bold_line_if_useful(self, lines: List[BoldDocumentLine], text_parts: List[str], has_bold_text: bool):
        clean_text = self._clean_ai2_line_text("".join(text_parts))
        if clean_text:
            lines.append(BoldDocumentLine(text=clean_text, has_bold_text=has_bold_text))

    def _clean_ai2_line_text(self, text: str) -> str:
        cleaned = text.replace("\xa0", " ")
        return re.sub(r"\s+", " ", cleaned).strip()

    def _should_skip_ai2_line(self, text: str) -> bool:
        cleaned = text.strip().lower()
        return not cleaned or cleaned.startswith("classification:")

    def _record_has_any_option(self, record: QuestionRecord) -> bool:
        return bool(record.option_text_1.strip() or record.option_text_2.strip() or 
                    record.option_text_3.strip() or record.option_text_4.strip())

    # ==================================================================
    # Format Option 2: AI Generated & Validated Functional Parser
    # ==================================================================
    def parse_ai_validated_file(self, file_path: Path) -> List[QuestionRecord]:
        lines = self._extract_document_lines_with_styles(file_path)
        records: List[QuestionRecord] = []
        current_record: Optional[QuestionRecord] = None
        current_question_title = ""
        expecting_question_stem = False

        for line in lines:
            text = line.text
            item_match = self._match_ai_item_start(text)

            if item_match:
                if current_record is not None:
                    records.append(current_record)

                current_record = QuestionRecord(
                    source_type="AI Generated & Validated",
                    source_document=file_path.name,
                    item_number=item_match.group("number").strip(),
                    question_text="",
                    question_title=current_question_title
                )
                expecting_question_stem = True
                continue

            if current_record is None:
                if line.is_heading or self._looks_like_ai_question_title(text):
                    current_question_title = text
                continue

            if expecting_question_stem:
                current_record.question_text = text
                expecting_question_stem = False
                continue

            option_match = self._match_option_line(text)
            if option_match:
                self._apply_option_to_record(current_record, option_match)
                continue

            answer_match = self._match_correct_answer(text)
            if answer_match:
                self._apply_correct_answer_to_record(current_record, answer_match)
                continue

            difficulty_match = self._match_metadata_line(text, "Difficulty")
            if difficulty_match:
                current_record.difficulty = difficulty_match
                continue

            rationale_match = self._match_metadata_line(text, "Rationale")
            if rationale_match:
                current_record.explanation = rationale_match
                continue

            if line.is_heading or self._looks_like_ai_question_title(text):
                if self._record_has_core_content(current_record):
                    records.append(current_record)
                    current_record = None
                    current_question_title = text
                    expecting_question_stem = False
                continue

        if current_record is not None:
            records.append(current_record)

        return self._filter_complete_records(records)

    def _extract_document_lines_with_styles(self, file_path: Path) -> List[DocumentLine]:
        document = Document(str(file_path))
        lines: List[DocumentLine] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue

            style_name = ""
            if paragraph.style is not None and paragraph.style.name is not None:
                style_name = paragraph.style.name.lower()
            is_heading = "heading" in style_name

            for raw_line in paragraph_text.splitlines():
                clean_line = raw_line.strip()
                if clean_line:
                    lines.append(DocumentLine(text=clean_line, is_heading=is_heading))
        return lines

    def _looks_like_ai_question_title(self, text: str) -> bool:
        if not text or len(text) > 90 or text.endswith("?"):
            return False
        if self._match_ai_item_start(text) or self._match_option_line(text) or self._match_correct_answer(text):
            return False
        
        known_metadata_labels = ["difficulty", "rationale", "taxonomy", "traceability", "correct answer", "answer", "key"]
        lowered = text.lower().strip()
        return not any(lowered.startswith(label) for label in known_metadata_labels)

    def _match_ai_item_start(self, text: str) -> Optional[re.Match]:
        return re.match(r"^\s*Item\s+(?P<number>\d+)\s*$", text, flags=re.IGNORECASE)

    # ==================================================================
    # Format Option 1: CE Formatted Hierarchical Table Parser
    # ==================================================================
    def parse_ce_formatted_file(self, file_path: Path) -> List[QuestionRecord]:
        document = Document(str(file_path))
        records: List[QuestionRecord] = []
        tables = list(document.tables)

        for table_index, table in enumerate(tables):
            item_number = self._get_ce_item_number(table)
            if not item_number:
                continue

            record = self._parse_ce_item_table(table, item_number, file_path.name)
            if record is None:
                continue

            record.question_title = self._find_ce_question_title_after_item(tables, table_index + 1)
            records.append(record)

        return self._filter_complete_records(records)

    def _get_ce_item_number(self, table: Table) -> str:
        if not table.rows or not table.rows[0].cells:
            return ""
        first_cell_text = table.rows[0].cells[0].text.strip()
        match = re.search(r"\b(?P<number>\d{6})\b", first_cell_text)
        return match.group("number") if match else ""

    def _parse_ce_item_table(self, item_table: Table, item_number: str, source_doc: str) -> Optional[QuestionRecord]:
        first_cell = item_table.rows[0].cells[0]
        nested_tables = list(first_cell.tables)
        if len(nested_tables) < 2:
            return None

        question_text = self._extract_ce_question_stem(nested_tables[0])
        options = self._extract_ce_options(nested_tables[1])
        if not question_text:
            return None

        record = QuestionRecord(
            source_type="CE Formatted",
            source_document=source_doc,
            item_number=item_number,
            question_text=question_text
        )
        record.option_text_1 = options.get("A", "")
        record.option_text_2 = options.get("B", "")
        record.option_text_3 = options.get("C", "")
        record.option_text_4 = options.get("D", "")
        record.option_correct = ANSWER_LETTER_TO_NUMBER.get(options.get("correct", ""), "")
        return record

    def _extract_ce_question_stem(self, stem_table: Table) -> str:
        for row in stem_table.rows:
            cells = row.cells
            for cell_index, cell in enumerate(cells):
                cell_text = self._clean_cell_text(cell.text)
                if re.match(r"^\d+\.$", cell_text) and cell_index + 1 < len(cells):
                    return self._clean_cell_text(cells[cell_index + 1].text)
            if len(cells) >= 2:
                first_text = self._clean_cell_text(cells[0].text)
                second_text = self._clean_cell_text(cells[1].text)
                if re.match(r"^\d+\.$", first_text) and second_text:
                    return second_text
        return ""

    def _extract_ce_options(self, options_table: Table) -> Dict[str, str]:
        options = {"A": "", "B": "", "C": "", "D": "", "correct": ""}
        for row in options_table.rows:
            row_cells = row.cells
            row_texts = [self._clean_cell_text(cell.text) for cell in row_cells]
            option_letter = self._find_option_letter_in_row(row_texts)
            if not option_letter:
                continue

            response_text = self._find_response_text_in_row(row_texts, option_letter)
            if response_text:
                options[option_letter] = response_text
            if any("*" in text for text in row_texts):
                options["correct"] = option_letter
        return options

    def _find_option_letter_in_row(self, row_texts: List[str]) -> str:
        for text in row_texts:
            match = re.match(r"^\s*(?P<letter>[A-Da-d])[\.\)]?\s*$", text)
            if match:
                return match.group("letter").upper()
        return ""

    def _find_response_text_in_row(self, row_texts: List[str], option_letter: str) -> str:
        found_letter = False
        for text in row_texts:
            if not text or ("*" in text and len(text.replace("*", "").strip()) == 0):
                continue
            if re.match(rf"^\s*{re.escape(option_letter)}[\.\)]?\s*$", text, flags=re.IGNORECASE):
                found_letter = True
                continue
            if found_letter:
                return text
        return ""

    def _find_ce_question_title_after_item(self, tables: List[Table], start_index: int) -> str:
        for table in tables[start_index:]:
            if self._get_ce_item_number(table):
                break
            title = self._extract_co_ref_from_table(table)
            if title:
                return title
        return ""

    def _extract_co_ref_from_table(self, table: Table) -> str:
        for row in table.rows:
            cells = row.cells
            for cell_index, cell in enumerate(cells):
                cell_text = self._clean_cell_text(cell.text)
                if self._is_co_ref_label(cell_text):
                    for next_cell in cells[cell_index + 1:]:
                        value = self._clean_cell_text(next_cell.text)
                        if value:
                            return value
        return ""

    def _is_co_ref_label(self, text: str) -> bool:
        cleaned = text.strip().lower().replace(":", "")
        return re.sub(r"\s+", " ", cleaned) == "co ref #"

    # ==================================================================
    # Shared Parsing & Evaluation Helpers
    # ==================================================================
    def _match_option_line(self, text: str) -> Optional[re.Match]:
        return re.match(r"^\s*(?P<letter>[A-Da-d])[\.\)\:\-]\s*(?P<text>.+)$", text)

    def _match_correct_answer(self, text: str) -> Optional[re.Match]:
        return re.match(r"^\s*(?:Correct\s+Answer|Answer|Key)\s*[\:\-]\s*(?P<answer>[A-Da-d])\b", text, flags=re.IGNORECASE)

    def _match_metadata_line(self, text: str, label: str) -> str:
        pattern = rf"^\s*{re.escape(label)}\s*[\:\-]\s*(?P<value>.+)$"
        match = re.match(pattern, text, flags=re.IGNORECASE)
        return match.group("value").strip() if match else ""

    def _apply_option_to_record(self, record: QuestionRecord, option_match: re.Match):
        option_letter = option_match.group("letter").upper()
        option_text = option_match.group("text").strip()
        if option_letter == "A": record.option_text_1 = option_text
        elif option_letter == "B": record.option_text_2 = option_text
        elif option_letter == "C": record.option_text_3 = option_text
        elif option_letter == "D": record.option_text_4 = option_text

    def _apply_correct_answer_to_record(self, record: QuestionRecord, answer_match: re.Match):
        ans_letter = answer_match.group("answer").strip().upper()
        record.option_correct = ANSWER_LETTER_TO_NUMBER.get(ans_letter, "")

    def _record_has_core_content(self, record: QuestionRecord) -> bool:
        return bool(record.question_text.strip() and record.option_text_1.strip() and 
                    record.option_text_2.strip() and record.option_text_3.strip() and record.option_text_4.strip())

    def _filter_complete_records(self, records: List[QuestionRecord]) -> List[QuestionRecord]:
        return [r for r in records if r.question_text.strip() and r.option_text_1.strip() and 
                r.option_text_2.strip() and r.option_text_3.strip() and r.option_text_4.strip()]

    def _clean_cell_text(self, text: str) -> str:
        cleaned = text.replace("\r", "\n").replace("\xa0", " ")
        parts = [p.strip() for p in cleaned.splitlines() if p.strip()]
        return " ".join(parts).strip()

    # ==================================================================
    # Optimized Pure-Python openpyxl Export Engine
    # ==================================================================
    def write_records_to_excel(self, records: List[QuestionRecord]) -> Path:
        """Streams records into a generated duplicate copy of the locked template asset via openpyxl."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Build clean filename appending current operational metadata
        user_name = re.sub(r'[<>:"/\\|?*]', "", getpass.getuser().strip()) or "User"
        date_text = datetime.now().strftime("%m-%d-%Y")
        base_name = f"{user_name}-Question Importer - {date_text}"
        
        suffix = self.template_path.suffix
        output_file = self.output_dir / f"{base_name}{suffix}"
        
        counter = 2
        while output_file.exists():
            output_file = self.output_dir / f"{base_name} ({counter}){suffix}"
            counter += 1

        # Duplicate locked template asset to workspace target output path
        shutil.copy2(self.template_path, output_file)

        # Initialize safe, in-memory workbook stream mapping
        wb = openpyxl.load_workbook(output_file)
        ws = wb.worksheets[0]

        # Dynamically scan the top rows to determine where data input should start
        first_data_row = 2
        for row in range(1, 31):
            cell_val = ws.cell(row=row, column=3).value
            if cell_val and str(cell_val).strip().lower().replace(" ", "") == "questiontext":
                first_data_row = row + 1
                break

        # Populate rows using exact structural parameters matching Column B through H
        for index, record in enumerate(records):
            current_row = first_data_row + index
            ws.cell(row=current_row, column=2, value=record.question_title)
            ws.cell(row=current_row, column=3, value=record.question_text)
            ws.cell(row=current_row, column=4, value=record.option_text_1)
            ws.cell(row=current_row, column=5, value=record.option_text_2)
            ws.cell(row=current_row, column=6, value=record.option_text_3)
            ws.cell(row=current_row, column=7, value=record.option_text_4)
            ws.cell(row=current_row, column=8, value=int(record.option_correct) if record.option_correct else "")

        wb.save(output_file)
        wb.close()
        
        self.log(f"Spreadsheet stream closed successfully. Document compiled at: {output_file.name}")
        return output_file

if __name__ == "__main__":
    import sys as _sys

    if len(_sys.argv) < 2:
        print("[EXAM-TRANSFORMER] Usage: python word_to_excel_exam.py <source_file_path> [ai|ai2|ce]")
        _sys.exit(1)

    source_path = Path(_sys.argv[1])
    format_type = _sys.argv[2] if len(_sys.argv) > 2 else "ai"

    engine = QuestionImporterEngine()
    try:
        result_path = engine.execute_pipeline(source_path, format_type)
    except Exception as err:
        print(f"[EXAM-TRANSFORMER] Pipeline failed: {err}")
        _sys.exit(1)

    if result_path is None:
        print("[EXAM-TRANSFORMER] No question records were extracted from the source document.")
        _sys.exit(1)

    print(f"[EXAM-TRANSFORMER] Output written to: {result_path}")