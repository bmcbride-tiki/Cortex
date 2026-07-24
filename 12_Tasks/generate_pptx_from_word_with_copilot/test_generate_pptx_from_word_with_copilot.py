# =============================================================================
# test_generate_pptx_from_word_with_copilot.py
# -----------------------------------------------------------------------------
# WHAT THIS FILE DOES
#   Checks that generate_pptx_from_word_with_copilot.py's `run()` produces
#   a real .pptx file on disk from a small generated .docx source, using
#   m365_graph_bridge's existing mock conversion logic.
#
# WHAT IT INTERACTS WITH
#   - `generate_pptx_from_word_with_copilot.py`, the file under test.
#   - `python-docx`, used here only to generate a throwaway source .docx
#     for the test to convert.
# =============================================================================

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from generate_pptx_from_word_with_copilot import GeneratePptxFromWordWithCopilot


def test_run_generates_pptx():
    with tempfile.TemporaryDirectory() as tmp:
        from docx import Document
        src = Path(tmp) / "source.docx"
        doc = Document()
        doc.add_paragraph("Section 1: Introduction")
        doc.add_paragraph("Body text.")
        doc.save(str(src))

        result = GeneratePptxFromWordWithCopilot().run(str(src), tmp, "")
        assert result["success"] is True
        assert Path(result["file_path"]).exists()


if __name__ == "__main__":
    test_run_generates_pptx()
    print("All generate_pptx_from_word_with_copilot self-checks passed.")
