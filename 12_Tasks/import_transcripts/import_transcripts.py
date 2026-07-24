import sys
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any

CURRENT_DIR = Path(__file__).resolve().parent
PARENT_DIR = CURRENT_DIR.parents[1]
CORE_DIR = PARENT_DIR / "00_System"

if str(CORE_DIR) not in sys.path:
    sys.path.append(str(CORE_DIR))

from database import get_db_connection

class TranscriptsImporter:
    def __init__(self):
        self.root_dir = PARENT_DIR
        self.inbox_dir = self._resolve_dir_case_insensitive(self.root_dir / "01_inbox" / "transcripts")
        self.vault_dir = self.root_dir / "02_vault" / "transcripts"

    def _resolve_dir_case_insensitive(self, target_path: Path) -> Path:
        """Helper to find the physical path even if it's named 01_Inbox instead of 01_inbox."""
        if target_path.exists():
            return target_path
        
        parent = target_path.parent
        if parent.exists():
            for child in parent.iterdir():
                if child.name.lower() == target_path.name.lower():
                    return child
        return target_path

    def _parse_filename(self, filename: str) -> Dict[str, str]:
        """
        Parses metadata from filename using: YYYY-MM-DD_Type_Trade_Title.[ext]
        """
        stem = Path(filename).stem
        parts = stem.split('_')
        
        metadata = {
            "date": datetime.now().strftime("%Y-%m-%d"),
            "type": "Internal",
            "trade": "General",
            "title": stem.replace("-", " ")
        }

        # Check for matching YYYY-MM-DD pattern
        if len(parts) >= 1:
            try:
                datetime.strptime(parts[0], "%Y-%m-%d")
                metadata["date"] = parts[0]
            except ValueError:
                pass

        # Check for standard 4-part name segment format
        if len(parts) >= 4:
            metadata["type"] = parts[1].strip()
            metadata["trade"] = parts[2].strip()
            metadata["title"] = " ".join(parts[3:]).replace("-", " ").strip()
        elif len(parts) == 3:
            metadata["type"] = parts[1].strip()
            metadata["title"] = parts[2].replace("-", " ").strip()

        return metadata

    def _read_docx(self, file_path: Path) -> str:
        """Robustly extracts text from both standard paragraphs and structured tables inside a DOCX file."""
        try:
            import docx
        except ImportError:
            raise ImportError("The 'python-docx' library is missing. Please run 'pip install python-docx' in your terminal.")

        doc = docx.Document(str(file_path))
        full_text = []

        # 1. Grab standard paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Grab table contents (Teams transcripts are often structured as Speaker | Time | Message tables)
        for table in doc.tables:
            for row in table.rows:
                # Merge duplicate cell text references (handling merged cell anomalies)
                seen_cells = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and (not seen_cells or seen_cells[-1] != text):
                        seen_cells.append(text)
                
                if seen_cells:
                    full_text.append(" | ".join(seen_cells))

        return "\n".join(full_text)

    def run(self) -> Dict[str, Any]:
        report = {"success": True, "files_processed": 0, "records_loaded": 0, "errors": []}

        # Assert directory setups
        self.inbox_dir.mkdir(parents=True, exist_ok=True)
        self.vault_dir.mkdir(parents=True, exist_ok=True)

        # Match both .txt and .docx extensions robustly (case-insensitive)
        supported_extensions = {".txt", ".docx"}
        files = [
            f for f in self.inbox_dir.iterdir() 
            if f.suffix.lower() in supported_extensions and not f.name.startswith("~$")
        ]

        if not files:
            report["errors"].append(f"No pending transcript files (.txt or .docx) found inside: {self.inbox_dir}")
            report["success"] = False
            return report

        conn = get_db_connection()
        cursor = conn.cursor()

        for file_path in files:
            try:
                # Extract clean body text based on file type
                if file_path.suffix.lower() == ".docx":
                    body_text = self._read_docx(file_path)
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        body_text = f.read().strip()

                if not body_text:
                    report["errors"].append(f"Empty file ignored: {file_path.name}")
                    continue

                # Parse metadata from filename
                meta = self._parse_filename(file_path.name)

                # 1. Archive the original file inside the vault (preserving original DOCX/TXT structure)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                vault_filename = f"{file_path.stem}_{timestamp}{file_path.suffix}"
                vault_path = self.vault_dir / vault_filename
                
                shutil.copy2(str(file_path), str(vault_path))

                # 2. Insert metadata record in catalog
                cursor.execute("""
                    INSERT OR REPLACE INTO transcripts_metadata (
                        title, meeting_date, source_type, trade, file_path, summary
                    ) VALUES (?, ?, ?, ?, ?, ?);
                """, (
                    meta["title"], meta["date"], meta["type"], 
                    meta["trade"], str(vault_path), 
                    f"Auto-imported {file_path.suffix.upper()} file. Total characters indexed: {len(body_text)}."
                ))
                origin_id = cursor.lastrowid

                # 3. Insert and index the raw text inside the Search Engine Index
                cursor.execute("""
                    INSERT INTO global_search_index (origin_id, content_type, body_text)
                    VALUES (?, 'transcript', ?);
                """, (origin_id, body_text))

                conn.commit()

                # Clean file out of inbox once committed
                file_path.unlink()

                report["files_processed"] += 1
                report["records_loaded"] += 1

            except Exception as e:
                conn.rollback()
                report["errors"].append(f"Error processing {file_path.name}: {str(e)}")
                report["success"] = False

        conn.close()
        return report

if __name__ == "__main__":
    importer = TranscriptsImporter()
    print(importer.run())